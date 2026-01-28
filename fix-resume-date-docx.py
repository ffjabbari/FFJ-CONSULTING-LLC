#!/usr/bin/env python3
"""
One-off: Replace "Mar 2017 – Nov 2017" with "Jan 2017 – Jan 2018" in the resume .docx.
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--user", "--quiet"])
    from docx import Document

DOCX_PATH = "/Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/Fred_Jabbari_Resume_Optimized_2026_with_Bedrock_BDAGood001_with_links.docx"
OLD_TEXT = "Mar 2017 – Nov 2017"
NEW_TEXT = "Jan 2017 – Jan 2018"


def replace_in_paragraph(paragraph):
    if OLD_TEXT not in paragraph.text:
        return False
    # Text can be split across runs; replace in full paragraph text then reassign
    new_text = paragraph.text.replace(OLD_TEXT, NEW_TEXT)
    if new_text == paragraph.text:
        return False
    # Clear and set first run to preserve formatting
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""
    return True


def main():
    p = Path(DOCX_PATH)
    if not p.exists():
        print(f"File not found: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(DOCX_PATH)
    changed = 0

    for para in doc.paragraphs:
        if replace_in_paragraph(para):
            changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para):
                        changed += 1

    if changed == 0:
        print("No occurrences of the date range found (may already be updated).")
    else:
        print(f"Replaced {changed} occurrence(s).")

    try:
        doc.save(DOCX_PATH)
        print(f"Saved: {DOCX_PATH}")
    except Exception as e:
        out = Path(__file__).parent / p.name
        doc.save(out)
        print(f"Could not write to original path: {e}")
        print(f"Saved to: {out}")


if __name__ == "__main__":
    main()
