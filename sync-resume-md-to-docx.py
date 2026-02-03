#!/usr/bin/env python3
"""
Sync the website resume (frontend/src/content/resume.md) to your local Word resume .docx.
Updates the .docx IN PLACE at your local path, or saves to repo if that path is not writable.

Run from repo root: python3 sync-resume-md-to-docx.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    import subprocess
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--user", "--quiet"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Paths
REPO_ROOT = Path(__file__).resolve().parent
MD_PATH = REPO_ROOT / "frontend" / "src" / "content" / "resume.md"
LOCAL_DOCX = Path("/Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/Fred_Jabbari_Resume_Optimized_2026_with_Bedrock_BDAGood001_with_links.docx")
FALLBACK_DOCX = REPO_ROOT / "Fred_Jabbari_Resume_synced_from_website.docx"


def add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink run to a paragraph."""
    try:
        from docx.oxml import parse_xml
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hl = parse_xml(
            f'<w:hyperlink r:id="{r_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r></w:hyperlink>'
        )
        paragraph._p.append(hl)
    except Exception:
        paragraph.add_run(text)


def add_inline_formatted(paragraph, line: str):
    """Add a line with **bold** and [text](url) support."""
    # Extract markdown links [text](url)
    rest = line
    while True:
        m = re.search(r'\[([^\]]+)\]\(([^)]+)\)', rest)
        if not m:
            break
        before = rest[:m.start()]
        link_text, url = m.group(1), m.group(2)
        rest = rest[m.end():]
        # Add bold segments in before
        for part in re.split(r'\*\*([^*]+)\*\*', before):
            if part:
                paragraph.add_run(part)
        add_hyperlink(paragraph, url, link_text)
    # Process remaining for **bold**
    for i, part in enumerate(re.split(r'\*\*([^*]+)\*\*', rest)):
        if not part:
            continue
        if i % 2 == 1:
            r = paragraph.add_run(part)
            r.bold = True
        else:
            # Handle *italic*
            for j, seg in enumerate(re.split(r'\*([^*]+)\*', part)):
                if not seg:
                    continue
                if j % 2 == 1:
                    run = paragraph.add_run(seg)
                    run.italic = True
                else:
                    paragraph.add_run(seg)


def md_to_docx(md_path: Path, out_path: Path) -> bool:
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()
    in_list = False

    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line
        stripped = line.strip()

        if stripped == "---":
            if in_list:
                in_list = False
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            if in_list:
                in_list = False
            p = doc.add_paragraph(stripped[2:].strip())
            p.style = "Title"
            i += 1
            continue

        if stripped.startswith("## "):
            if in_list:
                in_list = False
            p = doc.add_paragraph(stripped[3:].strip())
            p.style = "Heading 2"
            i += 1
            continue

        if stripped.startswith("### "):
            if in_list:
                in_list = False
            p = doc.add_paragraph(stripped[4:].strip())
            p.style = "Heading 3"
            i += 1
            continue

        if stripped.startswith("- "):
            if not in_list:
                in_list = True
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatted(p, bullet_text)
            i += 1
            continue

        if stripped.startswith("  - "):
            bullet_text = stripped[4:].strip()
            p = doc.add_paragraph(style="List Bullet 2")
            add_inline_formatted(p, bullet_text)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            if in_list:
                in_list = False
            p = doc.add_paragraph()
            add_inline_formatted(p, stripped[1:-1].strip())
            i += 1
            continue

        if stripped:
            if in_list:
                in_list = False
            p = doc.add_paragraph()
            add_inline_formatted(p, stripped)
        else:
            if in_list:
                in_list = False
            doc.add_paragraph()

        i += 1

    doc.save(out_path)
    return True


def main():
    print("=" * 60)
    print("Sync website resume (resume.md) → local Word (.docx)")
    print("=" * 60)

    if not MD_PATH.exists():
        print(f"❌ Resume markdown not found: {MD_PATH}")
        sys.exit(1)

    out_path = LOCAL_DOCX
    try:
        if not out_path.parent.exists():
            raise FileNotFoundError(f"Folder does not exist: {out_path.parent}")
        md_to_docx(MD_PATH, out_path)
        print(f"✅ Resume synced successfully (in place).")
        print(f"📄 Saved: {out_path}")
    except (PermissionError, FileNotFoundError) as e:
        out_path = FALLBACK_DOCX
        md_to_docx(MD_PATH, out_path)
        print(f"✅ Resume synced to repo.")
        print(f"📄 Saved: {out_path}")
        print(f"\nTo update your local resume in place, run:")
        print(f'  cp "{out_path}" "{LOCAL_DOCX}"')
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
