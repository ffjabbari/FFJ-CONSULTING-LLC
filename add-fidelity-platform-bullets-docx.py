#!/usr/bin/env python3
"""
Insert additional "platform leadership" bullets into the Fidelity (current assignment)
section of the Word resume (.docx).

Strategy:
- Find the paragraph that starts the next section: "Client: Nike" (or contains "Client: Nike")
- Insert the Fidelity bullets immediately BEFORE that paragraph
- Idempotent: if a unique bullet already exists, do nothing

Run locally on your Mac so it can write to the @@@PUBLIC path.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess

    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "--user", "--quiet"]
    )
    from docx import Document
    from docx.oxml import OxmlElement


DOCX_PATH = "/Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/Fred_Jabbari_Resume_Optimized_2026_with_Bedrock_BDAGood001_with_links.docx"

FIDELITY_MARKER = "Client: Fidelity"
NEXT_SECTION_MARKER = "Client: Nike"

UNIQUE_SENTINEL = "AWS-hosted Financial Planning platforms"

BULLETS = [
    "Led design of AWS-hosted Financial Planning platforms with a focus on scalability and accessibility.",
    "Defined modernization roadmaps integrating legacy systems with cloud-native services; partnered with stakeholders to drive alignment and adoption.",
    "Implemented IdP/SSO integrations (Okta) for authentication and authorization to AWS-based services.",
    "Delivered Infrastructure as Code using CloudFormation, Terraform, AWS CDK, and CDKTF; built POCs for automated AWS tagging with EventBridge.",
    "Built CI/CD pipelines using GitHub Actions, Jenkins, Octopus, and env0 for AWS and on‑prem deployments.",
    "Monitored and optimized systems using CloudWatch and Trusted Advisor to improve reliability and resource utilization.",
    "Applied PII and accessibility guidelines; implemented AWS security best practices (IAM, VPC, encryption) in regulated environments.",
]


def _insert_paragraph_before(paragraph, text: str):
    """Insert a new paragraph before `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph("")
    # Replace created paragraph xml with our inserted one
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para.add_run(text)
    return new_para


def _apply_list_style(paragraph, candidates):
    for style_name in candidates:
        try:
            paragraph.style = style_name
            return
        except Exception:
            continue


def main():
    path = Path(DOCX_PATH)
    if not path.exists():
        print(f"❌ File not found: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(DOCX_PATH)

    # Idempotency: if sentinel exists anywhere, assume already added.
    if any(UNIQUE_SENTINEL.lower() in (p.text or "").lower() for p in doc.paragraphs):
        print("✅ Fidelity platform bullets already present; no changes made.")
        return

    # Verify we can find Fidelity section at all (sanity)
    if not any(FIDELITY_MARKER.lower() in (p.text or "").lower() for p in doc.paragraphs):
        print(f"⚠️ Could not find '{FIDELITY_MARKER}' marker; proceeding anyway.")

    # Find the start of the next section (Nike) so we can insert just before it
    nike_para = None
    for p in doc.paragraphs:
        if NEXT_SECTION_MARKER.lower() in (p.text or "").lower():
            nike_para = p
            break

    if nike_para is None:
        print(f"❌ Could not find '{NEXT_SECTION_MARKER}' in the document. No changes made.")
        sys.exit(2)

    # Insert bullets before Nike section, preserving list look as best-effort
    # Insert in reverse order so final order matches BULLETS
    for bullet in reversed(BULLETS):
        bp = _insert_paragraph_before(nike_para, bullet)
        _apply_list_style(bp, ["List Bullet", "List Paragraph"])

    try:
        doc.save(DOCX_PATH)
        print(f"✅ Updated and saved in place: {DOCX_PATH}")
    except Exception as e:
        out = Path.cwd() / path.name.replace(".docx", "_updated.docx")
        doc.save(out)
        print(f"⚠️ Could not save in place: {e}")
        print(f"✅ Saved updated file to: {out}")


if __name__ == "__main__":
    main()

