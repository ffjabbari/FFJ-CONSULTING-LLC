#!/usr/bin/env python3
"""
Extract text content from .docx files
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Install with: pip3 install python-docx")
    sys.exit(1)

def extract_docx_text(file_path):
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    except Exception as e:
        return f"Error extracting content: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 extract_docx.py <docx_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    content = extract_docx_text(file_path)
    print(content)
