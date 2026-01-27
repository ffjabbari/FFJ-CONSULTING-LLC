#!/usr/bin/env python3
"""
Update the .docx resume file with website links
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--user", "--quiet"])
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

RESUME_PATH = "/Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/Fred_Jabbari_Resume_Optimized_2026_with_Bedrock_BDAGood001.docx"
WEBSITE_URL = "https://ffjconsultingllc.com"
GITHUB_URL = "https://github.com/ffjabbari/FFJ-CONSULTING-LLC"

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the hyperlink XML with proper namespace declarations
    hyperlink_xml = (
        '<w:hyperlink r:id="%s" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t>%s</w:t></w:r>'
        '</w:hyperlink>' % (r_id, text)
    )
    
    hyperlink = parse_xml(hyperlink_xml)
    paragraph._p.append(hyperlink)

def update_resume():
    """Add links section to the resume"""
    print(f"Opening resume: {RESUME_PATH}")
    
    if not Path(RESUME_PATH).exists():
        print(f"❌ Resume file not found: {RESUME_PATH}")
        return False
    
    try:
        doc = Document(RESUME_PATH)
        
        # Add a new section at the end
        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Empty line
        
        # Add heading
        heading = doc.add_paragraph("Additional Resources")
        heading.style = 'Heading 2'
        
        # Add website link
        p1 = doc.add_paragraph()
        p1.add_run("FFJ Consulting LLC Website: ").bold = True
        add_hyperlink(p1, WEBSITE_URL, WEBSITE_URL)
        
        # Add article link
        p2 = doc.add_paragraph()
        p2.add_run("AI History, Past and Present: ").bold = True
        article_url = f"{WEBSITE_URL}/article/ai-revolution-demo"
        add_hyperlink(p2, article_url, article_url)
        
        # Add GitHub link
        p3 = doc.add_paragraph()
        p3.add_run("GitHub Source Code: ").bold = True
        add_hyperlink(p3, GITHUB_URL, GITHUB_URL)
        
        # Save the updated document to current directory
        import os
        resume_filename = os.path.basename(RESUME_PATH)
        output_filename = resume_filename.replace('.docx', '_with_links.docx')
        output_path = os.path.join(os.getcwd(), output_filename)
        doc.save(output_path)
        
        print(f"✅ Resume updated successfully!")
        print(f"📄 Saved as: {output_path}")
        print(f"\nThe original file was preserved.")
        print(f"Review the new file and replace the original if you're happy with it.")
        return True
        
    except Exception as e:
        print(f"❌ Error updating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("="*60)
    print("Updating Resume with Website Links")
    print("="*60)
    print()
    
    if update_resume():
        print("\n✅ Done!")
    else:
        print("\n❌ Failed to update resume")
        sys.exit(1)
