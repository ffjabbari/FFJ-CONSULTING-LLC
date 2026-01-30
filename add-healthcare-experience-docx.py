#!/usr/bin/env python3
"""
Insert "General Healthcare Experience (Selected)" bullets under the SAIC entry
in the Word resume (.docx).

This script is designed to be run locally on your Mac so it can write to:
  /Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/...
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess

    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "--user", "--quiet"]
    )
    from docx import Document
    from docx.oxml import OxmlElement


DOCX_PATH = "/Users/fjabbari/@@@PUBLIC/@@@RESUME_2026/Fred_Jabbari_Resume_Optimized_2026_with_Bedrock_BDAGood001_with_links.docx"

HEADING_TEXT = "General Healthcare Experience (Selected)"
HEALTHCARE_BULLETS = [
    "Led design of AWS-hosted CMS platforms with a focus on scalability and accessibility.",
    "Defined modernization roadmaps integrating legacy systems with cloud-native services; partnered with stakeholders to drive alignment and adoption.",
    "Implemented IdP/SSO integrations (Okta) for authentication and authorization to AWS-based services.",
    "Delivered Infrastructure as Code using CloudFormation, Terraform, AWS CDK, and CDKTF; built POCs for automated AWS tagging with EventBridge.",
    "Built CI/CD pipelines using GitHub Actions, Jenkins, Octopus, and env0 for AWS and on‑prem deployments.",
    "Monitored and optimized systems using CloudWatch and Trusted Advisor to improve reliability and resource utilization.",
    "Applied PII and accessibility guidelines; ensured compliance with HIPAA and MARS‑E and AWS security best practices (IAM, VPC, encryption).",
]


def _insert_paragraph_after(paragraph, text: str):
    """Insert a new paragraph after `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph("")
    # Swap the created paragraph's XML with our inserted one
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para.add_run(text)
    return new_para


def _apply_list_style(paragraph, candidates):
    """Try list styles in order; ignore if missing."""
    for style_name in candidates:
        try:
            paragraph.style = style_name
            return
        except Exception:
            continue


def main():
    path = Path(DOCX_PATH)
    if not path.exists():
        print(f"❌ File not found: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(DOCX_PATH)

    # Idempotency: don't insert twice
    if any(HEADING_TEXT.lower() in (p.text or "").lower() for p in doc.paragraphs):
        print("✅ Heading already present; no changes made.")
        return

    # Find SAIC paragraph
    saic_para = None
    for p in doc.paragraphs:
        if "saic" in (p.text or "").lower():
            saic_para = p
            break

    if saic_para is None:
        print("❌ Could not find 'SAIC' in the document. No changes made.")
        sys.exit(2)

    # Insert heading as a sub-bullet-like line
    heading_para = _insert_paragraph_after(saic_para, HEADING_TEXT)
    # Try common list styles (may vary depending on the template)
    _apply_list_style(heading_para, ["List Bullet 2", "List Bullet"])
    # Make heading bold (best-effort)
    try:
        for run in heading_para.runs:
            run.bold = True
    except Exception:
        pass

    # Insert bullets under heading
    last = heading_para
    for bullet in HEALTHCARE_BULLETS:
        p = _insert_paragraph_after(last, bullet)
        _apply_list_style(p, ["List Bullet 3", "List Bullet 2", "List Bullet"])
        last = p

    # Save (prefer in-place; fallback to workspace/local directory)
    try:
        doc.save(DOCX_PATH)
        print(f"✅ Updated and saved in place: {DOCX_PATH}")
    except Exception as e:
        out = Path.cwd() / path.name.replace(".docx", "_updated.docx")
        doc.save(out)
        print(f"⚠️ Could not save in place: {e}")
        print(f"✅ Saved updated file to: {out}")


if __name__ == "__main__":
    main()

